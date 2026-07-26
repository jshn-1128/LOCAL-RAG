"""
DOCX file loader (.docx).

Handles Word documents using python-docx for text extraction and metadata.
"""

from __future__ import annotations

import logging
import os
from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument

from app.domain.exceptions import CorruptedDocumentError
from app.domain.models.document import Document, DocumentMetadata
from app.domain.ports.document_loader import DocumentLoaderPort
from app.infrastructure.document_loaders.checksum_generator import generate_checksum
from app.infrastructure.document_loaders.metadata_extractor import (
    extract_content_stats,
    extract_docx_metadata,
)
from app.infrastructure.document_loaders.mime_type_detector import detect_mime_type

logger = logging.getLogger(__name__)


class DocxLoader(DocumentLoaderPort):
    async def load(self, source: Path) -> Document:
        return self._load_file(source)

    async def load_many(self, sources: list[Path]) -> list[Document]:
        return [self._load_file(s) for s in sources]

    @property
    def supported_extensions(self) -> set[str]:
        return {".docx"}

    def _load_file(self, path: Path) -> Document:
        mime_type = detect_mime_type(".docx")

        try:
            docx = DocxDocument(str(path))
        except Exception as exc:
            raise CorruptedDocumentError(f"Failed to read DOCX: {path}. {exc}") from exc

        paragraphs = [p.text for p in docx.paragraphs]
        text = "\n".join(paragraphs)
        checksum = generate_checksum(text)
        file_meta = extract_content_stats(text)

        docx_meta = extract_docx_metadata(docx.core_properties)
        author = docx_meta.get("author")
        title = docx_meta.get("title") or path.stem

        stat = os.stat(str(path))

        return Document(
            content=text,
            source_path=path.resolve(),
            filename=path.name,
            title=title,
            checksum=checksum,
            file_type=".docx",
            mime_type=mime_type,
            encoding="utf-8",
            created_at=(
                datetime.fromtimestamp(stat.st_birthtime)
                if hasattr(stat, "st_birthtime")
                else None
            ),
            modified_at=datetime.fromtimestamp(stat.st_mtime),
            metadata=DocumentMetadata(
                author=author,
                word_count=file_meta.word_count,
                character_count=file_meta.character_count,
                custom=docx_meta,
            ),
        )
